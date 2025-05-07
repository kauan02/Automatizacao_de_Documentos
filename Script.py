import json
import os
import datetime
import xlwings as xw
from docx import Document

json_path = 'data.json'
modelos = {
    "LIO": './Documentos_Base/Lista_de_IO_Base.xlsx',
    "PTP": './Documentos_Base/Procedimento_de_TAF_do_Painel_Base.docx',
    "RTP": './Documentos_Base/Relatorio_de_TAF_do_Painel_Base.docx',
    "MCE": './Documentos_Base/Mapa_de_Comunicacao_Equipamento_Base.xlsx',
    "MCC": './Documentos_Base/Mapa_de_Comunicacao_CLP_Prosoft_Base.xlsx',
    "PTA": './Documentos_Base/Procedimento_TAF_Software_Aplicativo_Base.docx',
    "RTA": './Documentos_Base/Relatorio_TAF_Software_Aplicativo_Base.docx',
    "MSC": './Documentos_Base/Mapa_de_Comunicacao_SCADA_Base.xlsx',
    "IOM": './Documentos_Base/IOM_Base.docx',
    "PTC": './Documentos_Base/Procedimento_TAC_Base.docx',
    "RTC": './Documentos_Base/Relatorio_TAC_Base.docx'
}
saida_dir = 'gerados'
os.makedirs(saida_dir, exist_ok=True)

with open(json_path, 'r', encoding='utf-8') as f:
    dados = json.load(f)

projeto = input("Escolha o projeto (opções: OIS001, OIS002): ").upper()
base = input("Número da base (ex: 06): ").zfill(2)
arquivo = input("Número do arquivo (ex: 07): ").zfill(2)
revisao = input("Revisão (ex: 1): ")
execucao = input("Execução (ex: Kauan Barbosa -> K.B.): ").upper()
data = datetime.date.today()

try:
    info_base = dados[projeto][base]
    info_doc = info_base["documentos"][arquivo]
except KeyError:
    print("❌ Projeto, base ou arquivo não encontrados.")
    exit()

substituicoes = {
    "<BASE_TITULO>": info_base.get("nome_base_titulo", ""),
    "<BASE>": info_base.get("nome_base", ""),
    "<ESTADO>": info_base.get("estado", ""),
    "<OT/SS/CC>": info_base.get("OT/SS/CC", ""),
    "<CODIGO_DOCUMENTO>": info_doc.get("codigo", ""),
    "<REV>": revisao,
    "<PAINEL>": info_base.get("painel", ""),
    "<DATA>": str(data),
    "<EXECUCAO>": execucao
}

tipo_doc = info_doc["tipo"]
modelo_path = modelos.get(tipo_doc)

if not modelo_path:
    print(f"❌ Tipo de documento '{tipo_doc}' sem modelo associado.")
    exit()

ext = os.path.splitext(modelo_path)[1].lower()

if ext == ".xlsx":
    app = xw.App(visible=True)
    wb = app.books.open(modelo_path)

    for sheet in wb.sheets:
        used_range = sheet.used_range
        values = used_range.value

        for i, row in enumerate(values):
            for j, cell in enumerate(row):
                if isinstance(cell, str):
                    for marcador, valor in substituicoes.items():
                        if marcador in cell:
                            cell = cell.replace(marcador, valor)
                    values[i][j] = cell

        used_range.value = values

    nome_arquivo = f"{info_doc['codigo']}_{revisao}.xlsx"
    caminho_saida = os.path.join(saida_dir, nome_arquivo)
    wb.save(caminho_saida)
    wb.close()
    app.quit()

elif ext == ".docx":
    if not os.path.exists(modelo_path):
        print(f"❌ Arquivo de modelo não encontrado em: {modelo_path}")
        exit()
    else:
        print(f"📄 Modelo encontrado: {modelo_path}")

    doc = Document(modelo_path)

    def substituir_em_paragrafo(paragrafo, substituicoes):
        for run in paragrafo.runs:
            for marcador, valor in substituicoes.items():
                if marcador in run.text:
                    run.text = run.text.replace(marcador, valor)

    def substituir_em_tabela(tabela, substituicoes):
        for row in tabela.rows:
            for cell in row.cells:
                for paragrafo in cell.paragraphs:
                    substituir_em_paragrafo(paragrafo, substituicoes)


    for paragrafo in doc.paragraphs:
        substituir_em_paragrafo(paragrafo, substituicoes)


    for tabela in doc.tables:
        substituir_em_tabela(tabela, substituicoes)

    nome_arquivo = f"{info_doc['codigo']}_{revisao}.docx"
    caminho_saida = os.path.join(saida_dir, nome_arquivo)
    doc.save(caminho_saida)

else:
    print(f"❌ Tipo de arquivo não suportado: {ext}")
    exit()

print(f"✅ Documento gerado com sucesso: {caminho_saida}")
